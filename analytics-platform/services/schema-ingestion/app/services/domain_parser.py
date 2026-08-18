"""Document parsing and text chunking service for Domain Knowledge Base.

Supports PDF (.pdf), Word (.docx), CSV (.csv), and Excel (.xlsx) formats.
"""
from __future__ import annotations

import io
import os
import csv
import structlog
from typing import List, Dict, Any

log = structlog.get_logger(__name__)


def parse_document(file_bytes: bytes, file_name: str) -> str:
    """Extract raw plain text from document bytes based on file extension."""
    ext = file_name.split(".")[-1].lower() if "." in file_name else ""
    
    if ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return _parse_docx(file_bytes)
    elif ext == "csv":
        return _parse_csv(file_bytes)
    elif ext in ("xlsx", "xls"):
        return _parse_excel(file_bytes)
    else:
        # Fallback to UTF-8 text decode
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pypdf if available, with robust fallback."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for idx, page in enumerate(reader.pages):
            txt = page.extract_text() or ""
            if txt.strip():
                pages_text.append(f"--- Page {idx + 1} ---\n{txt}")
        return "\n\n".join(pages_text)
    except Exception as exc:
        log.warning("pdf_parsing_pypdf_fallback", error=str(exc))
        # Direct string fallback for raw text in pdf
        raw = file_bytes.decode("utf-8", errors="ignore")
        cleaned = "".join(c for c in raw if c.isprintable() or c in "\n\r\t")
        return cleaned[:50000]


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx if available."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_txt = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_txt:
                    paragraphs.append(row_txt)
        return "\n".join(paragraphs)
    except Exception as exc:
        log.warning("docx_parsing_fallback", error=str(exc))
        raw = file_bytes.decode("utf-8", errors="ignore")
        return "".join(c for c in raw if c.isprintable() or c in "\n\r\t")[:50000]


def _parse_csv(file_bytes: bytes) -> str:
    """Extract tabular summaries and header information from CSV."""
    text_content = file_bytes.decode("utf-8", errors="ignore")
    lines = text_content.splitlines()
    if not lines:
        return ""
    
    # Read headers and top sample rows
    reader = csv.reader(lines)
    rows = list(reader)
    if not rows:
        return ""

    headers = rows[0]
    output = [f"CSV Structure (Total Rows: {len(rows) - 1})", f"Columns: {', '.join(headers)}", "\nSample Data Rows:"]
    for row in rows[1:30]:  # Include top 30 rows
        output.append(" | ".join(row))
    return "\n".join(output)


def _parse_excel(file_bytes: bytes) -> str:
    """Extract text and tables from Excel workbook (.xlsx)."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        sheets_text = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheets_text.append(f"=== Sheet: {sheet_name} ===")
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(val).strip() for val in row if val is not None and str(val).strip()]
                if row_vals:
                    sheets_text.append(" | ".join(row_vals))
        return "\n".join(sheets_text)
    except Exception as exc:
        log.warning("excel_parsing_fallback", error=str(exc))
        return _parse_csv(file_bytes)


def chunk_text(text: str, max_chars: int = 1000, overlap: int = 100) -> List[str]:
    """Split extracted text into overlapping chunks for RAG embedding."""
    if not text or not text.strip():
        return []
    
    cleaned = text.strip()
    chunks = []
    start = 0
    length = len(cleaned)

    while start < length:
        end = start + max_chars
        if end < length:
            # Try to break at paragraph or line break
            break_idx = cleaned.rfind("\n", start, end)
            if break_idx == -1 or break_idx < start + (max_chars // 2):
                break_idx = cleaned.rfind(" ", start, end)
            if break_idx != -1 and break_idx > start:
                end = break_idx

        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap if (end - overlap) > start else end

    return chunks
